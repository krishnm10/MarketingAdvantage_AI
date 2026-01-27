# =============================================
# docx_parser_v2.py — Enhanced Word Document Parser (Production-Ready)
# Fully compatible with ingestion_v2 architecture
# Now includes unified LLM normalization toggle (global + local)
# =============================================

from typing import Dict, Any, List
from docx import Document
import os
import asyncio

from app.utils.text_cleaner_v2 import clean_text
from app.utils.logger import log_info, log_warning
from app.services.ingestion.llm_rewriter import rewrite_batch  # ✅ LLM integration
from app.config.ingestion_settings import ENABLE_LLM_NORMALIZATION  # ✅ Global toggle

# -------------------------------------------------------------------
# Local parser-level LLM toggle
# -------------------------------------------------------------------
# True  → Force enable LLM normalization for DOCX parser
# False → Force disable LLM normalization
# None  → Inherit from global ENABLE_LLM_NORMALIZATION
LOCAL_LLM_TOGGLE = None


def is_llm_enabled() -> bool:
    """Determine whether LLM normalization is enabled for this parser."""
    return ENABLE_LLM_NORMALIZATION if LOCAL_LLM_TOGGLE is None else LOCAL_LLM_TOGGLE


# -------------------------------------------------------------------
# TEXT EXTRACTORS
# -------------------------------------------------------------------
def extract_paragraphs(doc: Document) -> List[str]:
    """Extract visible paragraph text blocks."""
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def extract_tables(doc: Document) -> List[str]:
    """Extract readable rows from DOCX tables."""
    rows_text = []
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                rows_text.append(" | ".join(row_cells))
    return rows_text


def extract_headers_footers(doc: Document) -> List[str]:
    """Extract header/footer text from all sections (if any)."""
    headers, footers = [], []

    for section in doc.sections:
        if hasattr(section, "header") and section.header and section.header.paragraphs:
            for p in section.header.paragraphs:
                if p.text.strip():
                    headers.append(p.text.strip())

        if hasattr(section, "footer") and section.footer and section.footer.paragraphs:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    footers.append(p.text.strip())

    return headers + footers


# -------------------------------------------------------------------
# MERGER
# -------------------------------------------------------------------
def merge_blocks(blocks: List[str]) -> str:
    """Combine extracted sections into one clean text blob."""
    return "\n\n".join(blocks)


# -------------------------------------------------------------------
# MAIN PARSER PIPELINE
# -------------------------------------------------------------------
async def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    High-quality DOCX ingestion parser (async-safe).
    Steps:
      1. Extract paragraphs, tables, headers/footers
      2. Merge and clean all content
      3. (Optional) Normalize text with LLM
      4. Return ingestion_v2-ready dict
    """

    log_info(f"[docx_parser_v2] Reading DOCX: {file_path}")

    # <<< PATCH: load DOCX in a background thread to avoid blocking event loop >>>
    try:
        loop = asyncio.get_running_loop()
        doc = await loop.run_in_executor(None, lambda: Document(file_path))
    except Exception as e:
        raise ValueError(f"[docx_parser_v2] Failed to read DOCX: {file_path}: {e}")

    paragraphs = extract_paragraphs(doc)
    tables = extract_tables(doc)
    headers_footers = extract_headers_footers(doc)

    all_blocks = paragraphs + tables + headers_footers

    if not all_blocks:
        raise ValueError(f"[docx_parser_v2] No readable content in file: {file_path}")

    combined = merge_blocks(all_blocks)
    cleaned = clean_text(combined)
    normalized_text = cleaned

    # ----------------------------------------------------------------
    # ✅ Optional LLM normalization
    # ----------------------------------------------------------------
    if is_llm_enabled():
        try:
            log_info("[docx_parser_v2] Sending text for LLM normalization...")
            normalized_results = await rewrite_batch([cleaned])
            if normalized_results:
                normalized_text = normalized_results[0]
                log_info("[docx_parser_v2] ✅ LLM normalization complete.")
        except Exception as e:
            log_warning(f"[docx_parser_v2] ⚠️ LLM normalization failed: {e}")
    else:
        log_info("[docx_parser_v2] LLM normalization skipped (disabled).")

    # ----------------------------------------------------------------
    # ✅ Structured Output
    # ----------------------------------------------------------------
    return {
        "raw_text": combined,
        "cleaned_text": cleaned,
        "normalized_text": normalized_text,
        "blocks": len(all_blocks),
        "source_type": "docx",
        "metadata": {
            "file_name": os.path.basename(file_path),
            "blocks": len(all_blocks),
            "paragraphs": len(paragraphs),
            "table_rows": len(tables),
            "headers_footers": len(headers_footers),
            "parser": "docx_v2 (python-docx + LLM optional)",
            "llm_normalization": is_llm_enabled(),
        },
    }
